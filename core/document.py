"""Generazione del documento finale .docx dalla bozza IMRaD + forest plot."""
from __future__ import annotations

import base64
import io
import re

from docx import Document
from docx.shared import Inches


def _add_runs(paragraph, text: str) -> None:
    """Aggiunge il testo gestendo il grassetto **...**."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def build_docx(review_md: str, forest_png_b64: str | None = None,
               meta: dict | None = None) -> bytes:
    doc = Document()
    for raw in review_md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif stripped in ("---", "***", "___"):
            continue
        elif line.lstrip().startswith(("- ", "* ")):
            _add_runs(doc.add_paragraph(style="List Bullet"), line.lstrip()[2:])
        else:
            _add_runs(doc.add_paragraph(), line)

    if forest_png_b64:
        doc.add_heading("Figura — Forest plot (meta-analisi)", level=2)
        try:
            doc.add_picture(io.BytesIO(base64.b64decode(forest_png_b64)), width=Inches(6.0))
        except Exception:  # noqa: BLE001
            doc.add_paragraph("[forest plot non disponibile]")

    if meta:
        doc.add_heading("Nota tecnica", level=2)
        doc.add_paragraph(
            f"Differenza media aggregata {meta.get('pooled_md')} "
            f"(95% CI {meta.get('ci_low')}–{meta.get('ci_high')}), I²={meta.get('i2')}%, "
            f"modello: {meta.get('model', 'effetti casuali')}. "
            "I valori statistici provengono da un motore di calcolo dedicato; "
            "il testo è una bozza generata da AI da rivedere a cura degli autori."
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
